"""docx / xlsx / pdf / txt / md / csv 最小文件生成。"""
from __future__ import annotations

import csv
import json
import re
import site
import sys
import uuid
from datetime import datetime, timedelta, timezone
from io import BytesIO
from pathlib import Path
from typing import Any

from app.config import OSS_UPLOAD_PREFIX, UPLOADS_DIR, UPLOAD_META_DIR, oss_configured
from app.providers import oss as oss_provider
from app.models import AttachmentLite, ChatRequest, HistoryMessageItem
from app.services import chat_service, model_capability_service
from app.services.attachment_service import (
    build_attachment_public,
    extract_text_from_attachment_meta,
    read_attachment_meta_any,
    save_attachment_meta_db,
)
from app.services import ofox_gemini_search_service, ofox_responses_service
from app.services import rag_service
from app.storage import safe_json_dump


_DIRECT_EXPORT_HINTS = (
    "导出",
    "变成word",
    "变成 word",
    "生成word",
    "生成 word",
    "变成pdf",
    "变成 pdf",
    "生成pdf",
    "生成 pdf",
    "生成文档",
    "转为word",
    "转为 word",
    "帮我生成word",
    "帮我生成 word",
    "转为pdf",
    "转为 pdf",
    "帮我生成pdf",
    "帮我生成 pdf",
    "帮我生成文档",
    "转成word",
    "转成pdf",
    "导成word",
    "导成pdf",
    "整理成word",
    "整理成pdf",
    "生成txt",
    "生成 txt",
    "变成txt",
    "变成 txt",
    "转为txt",
    "转为 txt",
    "导出txt",
    "导出 txt",
    "变成markdown",
    "变成 markdown",
    "生成markdown",
    "生成 markdown",
    "转为markdown",
    "转为 markdown",
    "导出markdown",
    "导出 markdown",
    "变成md",
    "变成 md",
    "生成md",
    "转为md",
    "转为 md",
    "导出md",
)

_MODIFY_HINTS = (
    "加上",
    "补上",
    "改成",
    "改为",
    "调整",
    "优化",
    "更新",
    "重新",
    "再来",
    "继续",
)

_NEW_CONTENT_HINTS = (
    "帮我写",
    "给我写",
    "写一个",
    "写一篇",
    "写一段",
    "创作",
    "编写",
    "撰写",
    "小说",
    "故事",
    "诗",
    "文章",
    "文案",
    "邮件",
    "剧本",
    "科幻",
    "悬疑",
    "童话",
)

_EXPORT_CONTEXT_HINTS = (
    "上一轮",
    "上面的",
    "上述",
    "刚才",
    "原文",
    "内容",
    "这段",
    "这篇",
    "这份",
    "这段话",
)


def _history_items_from_task(task_data: dict) -> list[HistoryMessageItem]:
    items: list[HistoryMessageItem] = []
    for row in task_data.get("history_messages") or []:
        if not isinstance(row, dict):
            continue
        attachments = []
        for att in row.get("attachments") or []:
            if not isinstance(att, dict):
                continue
            attachments.append(
                AttachmentLite(
                    id=str(att.get("id") or ""),
                    name=str(att.get("name") or ""),
                    category=str(att.get("category") or "document"),
                )
            )
        items.append(
            HistoryMessageItem(
                role=str(row.get("role") or "user"),
                text=str(row.get("text") or ""),
                attachments=attachments,
                has_image_result=bool(row.get("has_image_result")),
            )
        )
    return items


def _latest_assistant_source_text(task_data: dict) -> str:
    for item in reversed(_history_items_from_task(task_data)):
        if item.role != "assistant":
            continue
        text = str(item.text or "").strip()
        if text and not _looks_like_generated_artifact_message(text):
            return text
    return ""


def _looks_like_generated_artifact_message(text: str) -> bool:
    normalized = str(text or "").strip().lower()
    if not normalized:
        return False
    if "[下载" in normalized or "](/api/attachments/" in normalized:
        return True
    return normalized.startswith("已生成word") or normalized.startswith("已生成pdf") or normalized.startswith("已生成excel") or normalized.startswith("已生成ppt") or normalized.startswith("已生成txt") or normalized.startswith("已生成markdown") or normalized.startswith("已生成csv")


def _latest_attachment_source(task_data: dict) -> tuple[str, str]:
    for attachment_id in reversed(list(task_data.get("attachment_ids") or [])):
        meta = read_attachment_meta_any(str(attachment_id))
        if not meta or meta.get("category") != "document":
            continue
        extracted = str(extract_text_from_attachment_meta(meta) or "").strip()
        if extracted:
            title = Path(str(meta.get("original_name") or "文档")).stem.strip() or "文档"
            return extracted, title
    return "", ""


def _best_export_source(task_data: dict) -> tuple[str, str]:
    attachment_text, attachment_title = _latest_attachment_source(task_data)
    if attachment_text:
        return attachment_text, attachment_title
    assistant_text = _latest_assistant_source_text(task_data)
    if assistant_text:
        return assistant_text, "文档"
    return "", ""


def _best_export_source_with_kind(task_data: dict) -> tuple[str, str, str]:
    attachment_text, attachment_title = _latest_attachment_source(task_data)
    if attachment_text:
        return attachment_text, attachment_title, "attachment"
    assistant_text = _latest_assistant_source_text(task_data)
    if assistant_text:
        return assistant_text, "文档", "assistant"
    return "", "", ""


def _direct_export_note_for_source(source_kind: str, artifact_type: str) -> str:
    target = {
        "docx": "Word 文档",
        "pdf": "PDF 文档",
        "txt": "TXT 文本文件",
        "md": "Markdown 文件",
    }.get(str(artifact_type or "").strip().lower(), "文档")
    if source_kind == "attachment":
        return f"已按当前附件原文直接转换为{target}，未重新概括，尽量避免缩短原文。"
    if source_kind == "assistant":
        return f"已按上一轮正文直接整理为{target}，未重新概括，尽量避免缩短原文。"
    return f"已按原始内容直接整理为{target}，尽量避免缩短原文。"


def _normalize_export_prompt(prompt: str) -> str:
    return re.sub(r"\s+", "", str(prompt or "").strip().lower())


def _looks_like_new_content_request(prompt: str) -> bool:
    normalized = _normalize_export_prompt(prompt)
    if not normalized:
        return False
    if re.search(r"\d+字", normalized):
        return True
    if any(token in normalized for token in _NEW_CONTENT_HINTS):
        if any(token in normalized for token in _EXPORT_CONTEXT_HINTS) and not any(
            token in normalized for token in ("小说", "故事", "诗", "文章", "文案", "邮件", "剧本", "科幻")
        ):
            return False
        return True
    return False


def _is_direct_export_followup(task_data: dict, artifact_type: str) -> bool:
    if artifact_type not in {"docx", "pdf", "txt", "md"}:
        return False
    source_text, _ = _best_export_source(task_data)
    if not source_text:
        return False
    prompt = _normalize_export_prompt(task_data.get("prompt") or "")
    if not prompt:
        return False
    if any(token in prompt for token in _MODIFY_HINTS):
        return False
    if _looks_like_new_content_request(prompt):
        return False
    if any(token in prompt for token in _DIRECT_EXPORT_HINTS):
        return True
    return len(prompt) <= 16 and any(token in prompt for token in ("word", "pdf", "文档", "txt", "markdown", "md"))


def _text_to_preserved_doc_schema(source_text: str, fallback_title: str = "文档") -> dict:
    raw = str(source_text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not raw:
        raise ValueError("缺少可导出的原始内容")

    lines = [line.rstrip() for line in raw.split("\n")]
    title = ""
    sections: list[dict[str, Any]] = []
    current_heading = "正文"
    current_paragraphs: list[str] = []

    def flush_section() -> None:
        nonlocal current_paragraphs
        paragraphs = [str(item).strip() for item in current_paragraphs if str(item).strip()]
        if paragraphs:
            sections.append({"heading": current_heading, "paragraphs": paragraphs})
        current_paragraphs = []

    heading_pattern = re.compile(r"^(#{1,6}\s*)?(\d+[\.\、]\s*)?.{1,40}[：:]?$")

    for idx, line in enumerate(lines):
        text = line.strip()
        if not text:
            if current_paragraphs and current_paragraphs[-1] != "":
                current_paragraphs.append("")
            continue

        if idx == 0 and len(text) <= 40 and not text.startswith(("-", "*", "•")):
            title = re.sub(r"^#{1,6}\s*", "", text).strip("：: ")
            continue

        compact = re.sub(r"^#{1,6}\s*", "", text).strip()
        if len(compact) <= 40 and heading_pattern.match(text) and not text.startswith(("-", "*", "•")):
            flush_section()
            current_heading = compact.strip("：: ")
            continue

        current_paragraphs.append(text)

    flush_section()
    if not sections:
        sections = [{"heading": "正文", "paragraphs": [raw]}]
    normalized_sections = []
    for section in sections:
        merged: list[str] = []
        buffer: list[str] = []
        for paragraph in section["paragraphs"]:
            if not paragraph:
                if buffer:
                    merged.append("\n".join(buffer).strip())
                    buffer = []
                continue
            buffer.append(paragraph)
        if buffer:
            merged.append("\n".join(buffer).strip())
        normalized_sections.append({
            "heading": section["heading"] or "正文",
            "paragraphs": merged or [raw],
        })
    return {"title": title or fallback_title, "sections": normalized_sections}


def _extract_json_object(text: str) -> dict:
    raw = str(text or "").strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
        raw = re.sub(r"\s*```$", "", raw)
    try:
        parsed = json.loads(raw)
        if isinstance(parsed, dict):
            return parsed
    except Exception:
        pass

    start = raw.find("{")
    end = raw.rfind("}")
    if start >= 0 and end > start:
        parsed = json.loads(raw[start : end + 1])
        if isinstance(parsed, dict):
            return parsed
    raise ValueError("模型没有返回可解析的 JSON 对象")


def _build_structured_messages(
    task_data: dict,
    artifact_type: str,
    *,
    knowledge_context: str = "",
    source_material: str = "",
) -> list[dict]:
    if artifact_type == "docx":
        system_prompt = (
            "你是 Word 文档结构生成器。"
            "请根据用户需求输出严格 JSON，不要解释，不要 Markdown 代码块。"
            '返回 schema: {"title": str, "sections": [{"heading": str, "paragraphs": [str]}]}.'
            "要求："
            "1. 标题简洁；"
            "2. 至少返回 1 个 section；"
            "3. paragraphs 只放纯文本段落；"
            "4. 不要包含样式、HTML、Markdown；"
            "5. 如果提供了待导出原文，除非用户明确要求压缩，否则尽量完整保留信息，不要擅自缩短。"
        )
    elif artifact_type == "xlsx":
        system_prompt = (
            "你是 Excel 结构生成器。"
            "请根据用户需求输出严格 JSON，不要解释，不要 Markdown 代码块。"
            '返回 schema: {"workbook_title": str, "subtitle": str, "data_note": str, "sheets": [{"name": str, "columns": [str], "rows": [[Any]]}]}.'
            "要求："
            "1. 至少返回 1 个 sheet；"
            "2. columns 为表头；"
            "3. rows 为二维数组；"
            "4. subtitle 和 data_note 可为空字符串；"
            "5. 不要包含公式、样式、HTML。"
        )
    elif artifact_type == "pdf":
        system_prompt = (
            "你是 PDF 文档结构生成器。"
            "请根据用户需求输出严格 JSON，不要解释，不要 Markdown 代码块。"
            '返回 schema: {"title": str, "sections": [{"heading": str, "paragraphs": [str]}]}.'
            "要求："
            "1. 标题简洁正式；"
            "2. 至少返回 1 个 section；"
            "3. paragraphs 只放纯文本段落；"
            "4. 适合打印阅读，不要包含样式、HTML、Markdown；"
            "5. 如果提供了待导出原文，除非用户明确要求压缩，否则尽量完整保留信息，不要擅自缩短。"
        )
    elif artifact_type in {"txt", "md"}:
        system_prompt = (
            "你是文档结构生成器。"
            "请根据用户需求输出严格 JSON，不要解释，不要 Markdown 代码块。"
            '返回 schema: {"title": str, "sections": [{"heading": str, "paragraphs": [str]}]}.'
            "要求："
            "1. 标题简洁；"
            "2. 至少返回 1 个 section；"
            "3. paragraphs 只放纯文本段落；"
            "4. 如果提供了待导出原文，除非用户明确要求压缩，否则尽量完整保留信息，不要擅自缩短。"
        )
    elif artifact_type == "csv":
        system_prompt = (
            "你是 CSV 结构生成器。"
            "请根据用户需求输出严格 JSON，不要解释，不要 Markdown 代码块。"
            '返回 schema: {"workbook_title": str, "subtitle": str, "data_note": str, "sheets": [{"name": str, "columns": [str], "rows": [[Any]]}]}.'
            "要求："
            "1. 至少返回 1 个 sheet；"
            "2. columns 为表头；"
            "3. rows 为二维数组；"
            "4. CSV 只会导出首个 sheet，所以首个 sheet 要最完整；"
            "5. 不要包含公式、样式、HTML。"
        )
    else:
        raise ValueError(f"暂不支持这种产物类型：{artifact_type}")

    current_user_content = chat_service.build_current_user_message_content(
        str(task_data.get("prompt") or ""),
        task_data.get("attachment_ids") or [],
    )
    if source_material.strip():
        guidance = (
            "\n\n以下是本轮需要整理进文件的待导出原文。"
            "除非用户明确要求压缩、摘要、提炼或改写，否则请尽量完整保留其中的关键信息，不要随意缩短。\n\n"
            f"【待导出原文】\n{source_material.strip()}"
        )
        if isinstance(current_user_content, str):
            current_user_content = (current_user_content + guidance).strip()
        elif isinstance(current_user_content, list):
            for block in current_user_content:
                if isinstance(block, dict) and block.get("type") == "text":
                    block["text"] = (str(block.get("text") or "") + guidance).strip()
                    break
    return chat_service.build_text_chat_messages(
        system_prompt=system_prompt,
        summary=str(task_data.get("summary") or ""),
        history_messages=_history_items_from_task(task_data),
        current_user_content=current_user_content,
        knowledge_context=knowledge_context,
    )


def _build_artifact_rag_bundle(task_data: dict) -> dict:
    if not bool(task_data.get("use_rag")):
        return {
            "used": False,
            "query": "",
            "sources": [],
            "context_text": "",
            "note": "",
            "rag_status": "",
        }
    req = ChatRequest(
        conversation_id=int(task_data.get("conversation_id") or 0),
        model=str(task_data.get("model") or ""),
        reasoning_mode=str(task_data.get("reasoning_mode") or "") or None,
        prompt=str(task_data.get("prompt") or ""),
        use_rag=bool(task_data.get("use_rag")),
        rag_query="",
        use_web_search=False,
        attachment_ids=list(task_data.get("attachment_ids") or []),
        summary=str(task_data.get("summary") or ""),
        history_messages=_history_items_from_task(task_data),
    )
    try:
        return rag_service.build_rag_bundle(req)
    except Exception:
        return {
            "used": False,
            "query": str(task_data.get("prompt") or "").strip(),
            "sources": [],
            "context_text": "",
            "note": "知识库服务暂不可用",
            "rag_status": "error",
        }


def _merge_artifact_sources(rag_sources: list[dict], web_sources: list[dict]) -> list[dict]:
    merged: list[dict] = []
    for row in rag_sources or []:
        if isinstance(row, dict):
            item = dict(row)
            item["source_type"] = "rag"
            merged.append(item)
    for row in web_sources or []:
        if isinstance(row, dict):
            item = dict(row)
            item["source_type"] = "web"
            merged.append(item)
    return merged


def _generate_structured_raw(task_data: dict, artifact_type: str) -> tuple[str, str, list[dict], str]:
    rag_bundle = _build_artifact_rag_bundle(task_data)
    source_material, _ = _latest_attachment_source(task_data)
    messages = _build_structured_messages(
        task_data,
        artifact_type,
        knowledge_context=rag_bundle.get("context_text") or "",
        source_material=source_material,
    )
    model = str(task_data.get("model") or "")
    adapter = model_capability_service.build_text_request_adapter(
        model,
        str(task_data.get("reasoning_mode") or "") or None,
    )
    note = chat_service.merge_response_notes(rag_bundle.get("note") or "")
    sources = _merge_artifact_sources(rag_bundle.get("sources") or [], [])
    rag_status = rag_bundle.get("rag_status") or ""

    if bool(task_data.get("use_web_search")) and chat_service.supports_builtin_web_search(model):
        try:
            parsed = ofox_responses_service.parse_responses_web_search_payload(
                ofox_responses_service.call_responses_api(
                    model=model,
                    input_payload=messages,
                    tools=ofox_responses_service.build_responses_web_search_tool(),
                    reasoning_effort=str(
                        (adapter.get("chat_completions_extra") or {}).get("reasoning_effort") or ""
                    ) or None,
                )
            )
            return (
                str(parsed.get("content") or ""),
                chat_service.merge_response_notes(note, "已使用模型内建联网搜索"),
                _merge_artifact_sources(rag_bundle.get("sources") or [], parsed.get("sources") or []),
                rag_status,
            )
        except Exception:
            note = chat_service.merge_response_notes(
                note,
                "模型内建联网搜索暂不可用，已按普通文件生成回答",
            )
    elif bool(task_data.get("use_web_search")) and chat_service.supports_gemini_builtin_web_search(model):
        try:
            parsed = ofox_gemini_search_service.parse_gemini_search_payload(
                ofox_gemini_search_service.call_gemini_search_api(
                    model=model,
                    input_payload=messages,
                    reasoning_mode=str(task_data.get("reasoning_mode") or "") or None,
                )
            )
            return (
                str(parsed.get("content") or ""),
                chat_service.merge_response_notes(note, "已使用模型内建联网搜索"),
                _merge_artifact_sources(rag_bundle.get("sources") or [], parsed.get("sources") or []),
                rag_status,
            )
        except Exception:
            note = chat_service.merge_response_notes(
                note,
                "模型内建联网搜索暂不可用，已按普通文件生成回答",
            )
    elif bool(task_data.get("use_web_search")):
        note = chat_service.merge_response_notes(
            note,
            chat_service.web_search_fallback_note(model),
        )

    raw = chat_service.call_chat_completion(
        model=model,
        messages=messages,
        reasoning_mode=str(task_data.get("reasoning_mode") or "") or None,
        temperature=0.2,
    )
    return raw, note, sources, rag_status


def _safe_file_stem(value: str, fallback: str) -> str:
    text = re.sub(r'[\\/:*?"<>|]+', "_", str(value or "").strip())
    text = re.sub(r"\s+", " ", text).strip().strip(".")
    if not text:
        text = fallback
    return text[:60]


def _save_generated_artifact(
    *,
    suffix: str,
    mime_type: str,
    display_name: str,
    data: bytes,
    task_id: str = "",
) -> dict:
    attachment_id = uuid.uuid4().hex
    stored_name = f"{attachment_id}{suffix}"
    stored_path = UPLOADS_DIR / stored_name
    storage = "local"
    oss_key = ""
    stored_path.write_bytes(data)
    if oss_configured():
        oss_key = f"{OSS_UPLOAD_PREFIX}/generated/{attachment_id}{suffix}"
        if oss_provider.write_bytes(oss_key, data, content_type=mime_type):
            storage = "oss"
    meta = {
        "id": attachment_id,
        "original_name": display_name,
        "stored_name": stored_name,
        "stored_path": str(stored_path),
        "storage": storage,
        "oss_key": oss_key if storage == "oss" else "",
        "suffix": suffix,
        "mime_type": mime_type,
        "category": "document",
        "size": len(data or b""),
    }
    safe_json_dump(UPLOAD_META_DIR / f"{attachment_id}.json", meta)
    save_attachment_meta_db(meta, task_id=task_id or None, source="generated_artifact")
    return build_attachment_public(meta)


def _normalize_docx_schema(data: dict) -> dict:
    title = str(data.get("title") or "文档").strip() or "文档"
    sections_raw = data.get("sections")
    sections: list[dict[str, Any]] = []
    if isinstance(sections_raw, list):
        for idx, row in enumerate(sections_raw):
            if not isinstance(row, dict):
                continue
            heading = str(row.get("heading") or f"第{idx + 1}部分").strip() or f"第{idx + 1}部分"
            paragraphs_raw = row.get("paragraphs")
            paragraphs = []
            if isinstance(paragraphs_raw, list):
                paragraphs = [str(item).strip() for item in paragraphs_raw if str(item).strip()]
            elif str(paragraphs_raw or "").strip():
                paragraphs = [str(paragraphs_raw).strip()]
            if not paragraphs:
                continue
            sections.append({"heading": heading, "paragraphs": paragraphs})
    if not sections:
        raise ValueError("文档结构为空，无法生成 Word")
    return {"title": title, "sections": sections}


def _normalize_xlsx_schema(data: dict) -> dict:
    workbook_title = str(data.get("workbook_title") or "数据表").strip() or "数据表"
    subtitle = str(data.get("subtitle") or "").strip()
    data_note = str(data.get("data_note") or "").strip()
    sheets_raw = data.get("sheets")
    sheets: list[dict[str, Any]] = []
    if isinstance(sheets_raw, list):
        for idx, row in enumerate(sheets_raw):
            if not isinstance(row, dict):
                continue
            name = str(row.get("name") or f"Sheet{idx + 1}").strip() or f"Sheet{idx + 1}"
            columns_raw = row.get("columns")
            columns = [str(item).strip() for item in columns_raw] if isinstance(columns_raw, list) else []
            columns = [item for item in columns if item]
            rows_raw = row.get("rows")
            rows: list[list[Any]] = []
            if isinstance(rows_raw, list):
                for item in rows_raw:
                    if isinstance(item, list):
                        rows.append(item)
            if not columns and rows:
                columns = [f"列{i + 1}" for i in range(len(rows[0]))]
            if not columns:
                continue
            sheets.append({"name": name[:31], "columns": columns, "rows": rows})
    if not sheets:
        raise ValueError("表格结构为空，无法生成 Excel")
    return {
        "workbook_title": workbook_title,
        "subtitle": subtitle,
        "data_note": data_note,
        "sheets": sheets,
    }


def _infer_xlsx_template(task_data: dict, schema: dict) -> str:
    text = " ".join(
        [
            str(task_data.get("prompt") or ""),
            str(schema.get("workbook_title") or ""),
            str(schema.get("subtitle") or ""),
            " ".join(str(sheet.get("name") or "") for sheet in schema.get("sheets") or []),
            " ".join(
                str(col)
                for sheet in schema.get("sheets") or []
                for col in (sheet.get("columns") or [])
            ),
        ]
    ).lower()
    report_tokens = (
        "行情",
        "汇总",
        "汇报",
        "报表",
        "分析",
        "趋势",
        "价格",
        "报价",
        "对比",
        "对照",
        "黄金",
        "统计",
    )
    if any(token in text for token in report_tokens):
        return "report_table"
    return "data_table"


def _render_docx_file(schema: dict) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(schema["title"], level=0)
    for section in schema["sections"]:
        doc.add_heading(section["heading"], level=1)
        for paragraph in section["paragraphs"]:
            doc.add_paragraph(str(paragraph))
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _render_txt_file(schema: dict) -> bytes:
    lines: list[str] = []
    title = str(schema.get("title") or "").strip()
    if title:
        lines.extend([title, "=" * max(3, min(len(title), 32)), ""])
    for section in schema.get("sections") or []:
        heading = str(section.get("heading") or "").strip()
        if heading:
            lines.extend([heading, "-" * max(3, min(len(heading), 24))])
        for paragraph in section.get("paragraphs") or []:
            lines.append(str(paragraph))
            lines.append("")
    return "\n".join(lines).strip().encode("utf-8")


def _render_md_file(schema: dict) -> bytes:
    lines: list[str] = []
    title = str(schema.get("title") or "").strip()
    if title:
        lines.append(f"# {title}")
        lines.append("")
    for section in schema.get("sections") or []:
        heading = str(section.get("heading") or "").strip()
        if heading:
            lines.append(f"## {heading}")
        for paragraph in section.get("paragraphs") or []:
            lines.append(str(paragraph))
            lines.append("")
    return "\n".join(lines).strip().encode("utf-8")


def _render_pdf_file(schema: dict) -> bytes:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfbase.pdfmetrics import registerFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        user_site = site.getusersitepackages()
        if user_site and user_site not in sys.path:
            sys.path.append(user_site)
            try:
                from reportlab.lib import colors
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
                from reportlab.lib.units import mm
                from reportlab.pdfbase.cidfonts import UnicodeCIDFont
                from reportlab.pdfbase.pdfmetrics import registerFont
                from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
            except ImportError as retry_exc:
                raise RuntimeError("PDF 依赖缺失，请先安装 reportlab。") from retry_exc
        else:
            raise RuntimeError("PDF 依赖缺失，请先安装 reportlab。") from exc

    registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "BaseCN",
        parent=styles["Normal"],
        fontName="STSong-Light",
        fontSize=10.5,
        leading=16,
        textColor=colors.HexColor("#1F2937"),
        spaceAfter=8,
    )
    title_style = ParagraphStyle(
        "TitleCN",
        parent=base,
        fontSize=18,
        leading=24,
        spaceAfter=16,
        textColor=colors.HexColor("#123B5D"),
    )
    heading_style = ParagraphStyle(
        "HeadingCN",
        parent=base,
        fontSize=13.5,
        leading=20,
        spaceBefore=8,
        spaceAfter=8,
        textColor=colors.HexColor("#0F172A"),
    )

    def _escape_pdf_text(text: str) -> str:
        value = str(text or "").strip()
        value = value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        return value.replace("\n", "<br/>")

    output = BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=schema["title"],
    )
    story = [Paragraph(_escape_pdf_text(schema["title"]), title_style), Spacer(1, 4)]
    for section in schema["sections"]:
        story.append(Paragraph(_escape_pdf_text(section["heading"]), heading_style))
        for paragraph in section["paragraphs"]:
            story.append(Paragraph(_escape_pdf_text(str(paragraph)), base))
        story.append(Spacer(1, 4))
    doc.build(story)
    return output.getvalue()


def _render_csv_file(schema: dict) -> bytes:
    sheets = schema.get("sheets") or []
    if not sheets:
        raise ValueError("CSV 结构为空，无法生成。")
    first = sheets[0]
    output = BytesIO()
    import io

    str_buffer = io.StringIO()
    writer = csv.writer(str_buffer)
    columns = list(first.get("columns") or [])
    if columns:
        writer.writerow(columns)
    for row in first.get("rows") or []:
        if isinstance(row, list):
            writer.writerow(row[: len(columns)] if columns else row)
    output.write(str_buffer.getvalue().encode("utf-8-sig"))
    return output.getvalue()


def _render_xlsx_file(schema: dict, template_type: str = "data_table") -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    wb.properties.creator = "怀仁AI中台"
    wb.properties.title = schema["workbook_title"]
    is_report = template_type == "report_table"
    title_fill = PatternFill("solid", fgColor="123B5D" if is_report else "1F4E78")
    subtitle_fill = PatternFill("solid", fgColor="EAF2F8" if is_report else "DCE6F1")
    header_fill = PatternFill("solid", fgColor="CFE2F3" if is_report else "D9EAF7")
    zebra_fill = PatternFill("solid", fgColor="F4F8FC" if is_report else "F7FBFF")
    thin_side = Side(style="thin", color="D0D7DE")
    border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)
    title_font = Font(name="Microsoft YaHei", size=18 if is_report else 16, bold=True, color="FFFFFF")
    subtitle_font = Font(name="Microsoft YaHei", size=10, color="3C4858")
    header_font = Font(name="Microsoft YaHei", size=11, bold=True, color="1F1F1F")
    body_font = Font(name="Microsoft YaHei", size=10, color="1F1F1F")
    note_font = Font(name="Microsoft YaHei", size=9, color="5B6573")
    sh_tz = timezone(timedelta(hours=8))
    generated_at = datetime.now(sh_tz).strftime("%Y-%m-%d %H:%M")

    def _is_empty(value: Any) -> bool:
        return value is None or str(value).strip() == ""

    def _coerce_value(value: Any) -> Any:
        if value is None:
            return ""
        if isinstance(value, (int, float, datetime)):
            return value
        raw = str(value).strip()
        if not raw:
            return ""
        normalized = raw.replace(",", "")
        if re.fullmatch(r"-?\d+", normalized):
            try:
                return int(normalized)
            except Exception:
                return raw
        if re.fullmatch(r"-?\d+\.\d+", normalized):
            try:
                return float(normalized)
            except Exception:
                return raw
        if re.fullmatch(r"\d{4}-\d{1,2}-\d{1,2}", raw):
            try:
                return datetime.strptime(raw, "%Y-%m-%d")
            except Exception:
                return raw
        if re.fullmatch(r"\d{4}/\d{1,2}/\d{1,2}", raw):
            try:
                return datetime.strptime(raw, "%Y/%m/%d")
            except Exception:
                return raw
        return raw

    def _column_role(column_name: str, values: list[Any]) -> str:
        name = str(column_name or "").strip().lower()
        if any(token in name for token in ("日期", "时间", "date", "day")):
            return "date"
        if any(token in name for token in ("金额", "价格", "单价", "总价", "预算", "收入", "支出", "费用", "cost", "amount", "price")):
            return "currency"
        if any(token in name for token in ("备注", "说明", "描述", "摘要", "内容", "comment", "note", "desc", "detail")):
            return "long_text"
        if any(token in name for token in ("名称", "项目", "品类", "主题", "title", "name")):
            return "wide_text"
        non_empty = [v for v in values if not _is_empty(v)]
        if non_empty and all(isinstance(v, datetime) for v in non_empty):
            return "date"
        if non_empty and all(isinstance(v, (int, float)) for v in non_empty):
            return "number"
        if any(len(str(v)) > 24 for v in non_empty):
            return "long_text"
        return "text"

    def _sheet_title(name: str, index: int) -> str:
        text = re.sub(r"[\\/*?:\[\]]+", "_", str(name or "").strip())
        text = text or f"报表{index + 1}"
        return text[:31]

    def _apply_column_widths(ws, widths: dict[int, int], roles: dict[int, str]) -> None:
        for idx, width in widths.items():
            letter = get_column_letter(idx)
            role = roles.get(idx, "text")
            if role == "long_text":
                final_width = max(18, min(width + 4, 42))
            elif role == "wide_text":
                final_width = max(14, min(width + 3, 28))
            elif role == "date":
                final_width = max(12, min(width + 2, 16))
            elif role in {"currency", "number"}:
                final_width = max(12, min(width + 2, 18))
            else:
                final_width = max(10, min(width + 2, 24))
            ws.column_dimensions[letter].width = final_width

    first = True
    for sheet_idx, sheet in enumerate(schema["sheets"]):
        if first:
            ws = wb.active
            ws.title = _sheet_title(sheet["name"] or schema["workbook_title"], sheet_idx)
            first = False
        else:
            ws = wb.create_sheet(title=_sheet_title(sheet["name"] or schema["workbook_title"], sheet_idx))

        columns = sheet["columns"]
        rows = sheet["rows"]
        col_count = max(1, len(columns))
        last_col_letter = get_column_letter(col_count)

        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=col_count)
        title_cell = ws.cell(row=1, column=1, value=schema["workbook_title"])
        title_cell.font = title_font
        title_cell.fill = title_fill
        title_cell.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[1].height = 32 if is_report else 28

        meta_line = schema.get("subtitle") or ("经营分析报表" if is_report else "正式数据台账")
        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=col_count)
        sub_cell = ws.cell(
            row=2,
            column=1,
            value=f"{meta_line}    生成时间：{generated_at}",
        )
        sub_cell.font = subtitle_font
        sub_cell.fill = subtitle_fill
        sub_cell.alignment = Alignment(horizontal="left", vertical="center")
        ws.row_dimensions[2].height = 20

        if schema.get("data_note"):
            ws.merge_cells(start_row=3, start_column=1, end_row=3, end_column=col_count)
            note_cell = ws.cell(row=3, column=1, value=f"数据说明：{schema['data_note']}")
            note_cell.font = note_font
            note_cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            ws.row_dimensions[3].height = 34
            header_row = 5
        else:
            header_row = 4

        for col_idx, column_name in enumerate(columns, start=1):
            cell = ws.cell(row=header_row, column=col_idx, value=column_name)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = border
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        ws.row_dimensions[header_row].height = 24

        column_values_map: dict[int, list[Any]] = {idx: [] for idx in range(1, col_count + 1)}
        for row in rows:
            padded = list(row) + [""] * max(0, col_count - len(row))
            for idx in range(1, col_count + 1):
                column_values_map[idx].append(_coerce_value(padded[idx - 1]))

        column_roles = {
            idx: _column_role(columns[idx - 1], values)
            for idx, values in column_values_map.items()
        }
        column_widths: dict[int, int] = {
            idx: (
                24
                if column_roles.get(idx) == "long_text"
                else 16
                if column_roles.get(idx) == "wide_text"
                else 12
                if column_roles.get(idx) == "date"
                else 14
                if column_roles.get(idx) in {"currency", "number"}
                else min(24, max(10, len(str(columns[idx - 1]))))
            )
            for idx in range(1, col_count + 1)
        }

        data_start_row = header_row + 1
        for row_offset, row in enumerate(rows, start=0):
            excel_row = data_start_row + row_offset
            padded = list(row) + [""] * max(0, col_count - len(row))
            max_line_count = 1
            for col_idx in range(1, col_count + 1):
                raw_value = _coerce_value(padded[col_idx - 1])
                cell = ws.cell(row=excel_row, column=col_idx, value=raw_value)
                cell.font = body_font
                cell.border = border
                if row_offset % 2 == 1:
                    cell.fill = zebra_fill

                role = column_roles[col_idx]
                horizontal = "left"
                vertical = "center"
                wrap_text = False
                if role in {"currency", "number"}:
                    horizontal = "right"
                if role == "date":
                    horizontal = "center"
                    if isinstance(raw_value, datetime):
                        cell.number_format = "yyyy-mm-dd"
                elif role == "currency":
                    if isinstance(raw_value, (int, float)):
                        cell.number_format = '#,##0.00'
                elif role == "number":
                    if isinstance(raw_value, (int, float)):
                        cell.number_format = '#,##0.###'
                elif role == "long_text":
                    wrap_text = True
                    vertical = "top"
                    if _is_empty(raw_value):
                        cell.value = ""
                elif role == "wide_text":
                    wrap_text = True
                    vertical = "top"
                if isinstance(raw_value, str):
                    lines = str(raw_value).count("\n") + 1
                    max_line_count = max(max_line_count, lines)
                    if role in {"long_text", "wide_text"} or len(raw_value) > 20:
                        wrap_text = True
                        vertical = "top"
                    display_len = max((len(seg) for seg in str(raw_value).splitlines()), default=0)
                    width_cap = 44 if role == "long_text" else 30 if role == "wide_text" else 24
                    width_guess = min(width_cap, max(len(str(columns[col_idx - 1])), display_len))
                    column_widths[col_idx] = max(column_widths[col_idx], width_guess)
                elif isinstance(raw_value, (int, float)):
                    column_widths[col_idx] = max(column_widths[col_idx], len(f"{raw_value:,.2f}"))
                elif isinstance(raw_value, datetime):
                    column_widths[col_idx] = max(column_widths[col_idx], 12)

                cell.alignment = Alignment(
                    horizontal=horizontal,
                    vertical=vertical,
                    wrap_text=wrap_text,
                )
            base_height = 20 if is_report else 18
            extra_height = 4 if any(column_roles.get(i) in {"long_text", "wide_text"} for i in range(1, col_count + 1)) else 0
            ws.row_dimensions[excel_row].height = min(base_height * max_line_count + extra_height, 84)

        ws.freeze_panes = f"A{data_start_row}"
        if rows:
            ws.auto_filter.ref = f"A{header_row}:{last_col_letter}{header_row + len(rows)}"
        _apply_column_widths(ws, column_widths, column_roles)

        for row_idx in range(1, data_start_row + len(rows)):
            for col_idx in range(1, col_count + 1):
                ws.cell(row=row_idx, column=col_idx).border = border
        ws.sheet_view.showGridLines = True

    output = BytesIO()
    wb.save(output)
    return output.getvalue()


def run_artifact_task(task_data: dict) -> dict:
    artifact_type = str(task_data.get("artifact_type") or "").strip().lower()
    if artifact_type == "pptx":
        from app.services import ppt_generation_service

        return ppt_generation_service.run_pptx_task(task_data)

    if artifact_type not in {"docx", "xlsx", "pdf", "txt", "md", "csv"}:
        raise RuntimeError(f"暂不支持这种文件任务：{artifact_type}")

    generation_sources: list[dict] = []
    rag_status = ""
    generation_note = ""
    template_type = ""

    if artifact_type in {"docx", "pdf", "txt", "md"} and _is_direct_export_followup(task_data, artifact_type):
        rag_bundle = _build_artifact_rag_bundle(task_data)
        source_text, source_title, source_kind = _best_export_source_with_kind(task_data)
        schema = _text_to_preserved_doc_schema(
            source_text,
            fallback_title=source_title or "文档",
        )
        generation_sources = _merge_artifact_sources(rag_bundle.get("sources") or [], [])
        rag_status = rag_bundle.get("rag_status") or ""
        generation_note = chat_service.merge_response_notes(
            rag_bundle.get("note") or "",
            _direct_export_note_for_source(source_kind, artifact_type),
        )
    else:
        raw, generation_note, generation_sources, rag_status = _generate_structured_raw(task_data, artifact_type)
        parsed = _extract_json_object(raw)
        schema = _normalize_docx_schema(parsed) if artifact_type in {"docx", "pdf", "txt", "md"} else None

    if artifact_type == "docx":
        if schema is None:
            raise RuntimeError("Word 文档结构为空。")
        filename = _safe_file_stem(schema["title"], "文档") + ".docx"
        content = _render_docx_file(schema)
        attachment = _save_generated_artifact(
            suffix=".docx",
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            display_name=filename,
            data=content,
            task_id=str(task_data.get("id") or ""),
        )
        label = "Word 文档"
    elif artifact_type == "txt":
        if schema is None:
            raise RuntimeError("TXT 文档结构为空。")
        filename = _safe_file_stem(schema["title"], "文档") + ".txt"
        content = _render_txt_file(schema)
        attachment = _save_generated_artifact(
            suffix=".txt",
            mime_type="text/plain; charset=utf-8",
            display_name=filename,
            data=content,
            task_id=str(task_data.get("id") or ""),
        )
        label = "TXT 文本文件"
    elif artifact_type == "md":
        if schema is None:
            raise RuntimeError("Markdown 文档结构为空。")
        filename = _safe_file_stem(schema["title"], "文档") + ".md"
        content = _render_md_file(schema)
        attachment = _save_generated_artifact(
            suffix=".md",
            mime_type="text/markdown; charset=utf-8",
            display_name=filename,
            data=content,
            task_id=str(task_data.get("id") or ""),
        )
        label = "Markdown 文件"
    elif artifact_type == "pdf":
        if schema is None:
            raise RuntimeError("PDF 文档结构为空。")
        filename = _safe_file_stem(schema["title"], "文档") + ".pdf"
        content = _render_pdf_file(schema)
        attachment = _save_generated_artifact(
            suffix=".pdf",
            mime_type="application/pdf",
            display_name=filename,
            data=content,
            task_id=str(task_data.get("id") or ""),
        )
        label = "PDF 文档"
    elif artifact_type == "csv":
        raw, generation_note, generation_sources, rag_status = _generate_structured_raw(task_data, artifact_type)
        parsed = _extract_json_object(raw)
        schema = _normalize_xlsx_schema(parsed)
        filename = _safe_file_stem(schema["workbook_title"], "数据表") + ".csv"
        content = _render_csv_file(schema)
        attachment = _save_generated_artifact(
            suffix=".csv",
            mime_type="text/csv; charset=utf-8",
            display_name=filename,
            data=content,
            task_id=str(task_data.get("id") or ""),
        )
        label = "CSV 文件"
    else:
        raw, generation_note, generation_sources, rag_status = _generate_structured_raw(task_data, artifact_type)
        parsed = _extract_json_object(raw)
        schema = _normalize_xlsx_schema(parsed)
        template_type = _infer_xlsx_template(task_data, schema)
        filename = _safe_file_stem(schema["workbook_title"], "数据表") + ".xlsx"
        content = _render_xlsx_file(schema, template_type=template_type)
        attachment = _save_generated_artifact(
            suffix=".xlsx",
            mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            display_name=filename,
            data=content,
            task_id=str(task_data.get("id") or ""),
        )
        label = "Excel 文件"

    link = f"/api/attachments/{attachment['id']}"
    return {
        "text": f"已生成{label}：[下载 {attachment['name']}]({link})",
        "note": (
            chat_service.merge_response_notes(
                generation_note,
                f"任务已完成，已按 {'报表型' if artifact_type == 'xlsx' and template_type == 'report_table' else '台账型' if artifact_type == 'xlsx' else ''}模板渲染。"
            )
            if artifact_type == "xlsx"
            else chat_service.merge_response_notes(generation_note, "任务已完成，已按 CSV 格式导出首个数据表。")
            if artifact_type == "csv"
            else chat_service.merge_response_notes(generation_note, "任务已完成，已按打印版文档渲染。")
            if artifact_type == "pdf"
            else chat_service.merge_response_notes(generation_note, "任务已完成，已按 Markdown 文本格式导出。")
            if artifact_type == "md"
            else chat_service.merge_response_notes(generation_note, "任务已完成，已按纯文本格式导出。")
            if artifact_type == "txt"
            else chat_service.merge_response_notes(generation_note, "任务已完成。")
        ),
        "artifact_type": artifact_type,
        "attachment": attachment,
        "sources": generation_sources,
        "rag_status": rag_status,
    }
